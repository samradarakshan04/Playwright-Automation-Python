from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font = run.font
    font.size = Pt(10.5)
    if bold:
        run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Header
add_heading("Samra Darakshan", level=0)
add_paragraph("Email: samra.darakshan@gmail.com | Phone: +92 335 2291683")
add_paragraph("LinkedIn: linkedin.com/in/samradarakshan | GitHub: github.com/yourprofile")

# Summary
add_heading("Professional Summary", level=1)
add_paragraph(
    "Experienced QA Automation Engineer with 7+ years of expertise in manual and automated testing across web, mobile, and chatbot platforms. "
    "Skilled in designing and maintaining automation frameworks using Selenium (Java/Python), Playwright, TestNG, PyTest, Cypress, Appium, and Cucumber. "
    "Proficient in API testing with Postman and Python, integrating tests into CI/CD pipelines using Jenkins and Git.\n\n"
    "Currently leading post-release UX auditing for eZhire’s customer app, driver app, and web portal — reviewing live UI, chatbot interactions, emails, and notifications to ensure consistency, clarity, and user satisfaction. Agile team player focused on quality, test coverage, and smooth user experiences."
)

# Skills
add_heading("Skills", level=1)
add_paragraph(
    "• Automation Tools: Selenium (Java/Python), Cypress, Playwright, Appium, Cucumber\n"
    "• API Testing: Postman, REST APIs, API Automation (JavaScript, Python)\n"
    "• Frameworks: PyTest, TestNG, BDD, POM\n"
    "• Programming: Python, Java, JavaScript\n"
    "• CI/CD: Jenkins, Docker\n"
    "• Version Control: Git, GitHub, Bitbucket\n"
    "• Soft Skills: Communication, Collaboration, Agile/Scrum, Problem Solving"
)

# Experience
add_heading("Professional Experience", level=1)

add_paragraph("Software QA Engineer | eZhire, Karachi — Nov 2024 – Present", bold=True)
add_paragraph(
    "- Built automation frameworks using Selenium (Python), Playwright, and Cypress.\n"
    "- Automated mobile apps with Appium (Python) and APIs using Postman.\n"
    "- Led post-release UX audits for customer, driver apps, and web portal.\n"
    "- Conducted manual testing, UAT support, and bug validation.\n"
    "- Collaborated via Git/GitHub and Jenkins in Agile environments."
)

add_paragraph("QA Automation Engineer | Royal Cyber, Karachi — Mar 2022 – Oct 2024", bold=True)
add_paragraph(
    "- Built web automation using Selenium (Java), Cypress, and API automation with Postman.\n"
    "- Integrated frameworks into Jenkins pipelines; versioned code in Bitbucket.\n"
    "- Automated admin APIs, performed load testing (JMeter), and chatbot QA (IBM Watson, Salesforce Copilot)."
)

add_paragraph("Software Engineer | Systems Ltd, Karachi — May 2020 – Mar 2022", bold=True)
add_paragraph(
    "- Maintained Angular/Strapi portals; manually tested UI and APIs.\n"
    "- Built REST APIs using IIB, ACE; collaborated with QA for end-to-end validation."
)

add_paragraph("Software Engineer | Royal Cyber, Karachi — Apr 2018 – May 2020", bold=True)
add_paragraph(
    "- Developed APIs with MuleSoft, React UI, Kafka integration.\n"
    "- Supported API QA, cloud migration, Docker, and Jenkins deployments."
)

# Key Projects
add_heading("Key Projects", level=1)

add_paragraph("UX Auditing – eZhire Customer App, Driver App & Web Portal", bold=True)
add_paragraph(
    "- Led post-release UX audits for all platforms, ensuring consistency and usability.\n"
    "- Reviewed live UI, chatbot flows, emails, and notifications for user experience issues."
)

add_paragraph("eZhire Automation Projects", bold=True)
add_paragraph(
    "- Built booking flow automation using Selenium (PyTest), and mobile automation using Appium.\n"
    "- API testing and automation for booking flows and vendor workflows via Postman (JavaScript/Python)."
)

add_paragraph("Trello & OrangeHRM Automation | Cypress POC", bold=True)
add_paragraph("- Automated regression cases using Selenium, Postman, Cypress for internal tools.")

# Education
add_heading("Education", level=1)
add_paragraph("MS Data Science | FAST-NUCES | Karachi | 2019 – 2021 | CGPA: 3.53", bold=True)
add_paragraph("BS Computer Science | Bahria University | Karachi | 2014 – 2018 | CGPA: 3.36", bold=True)

# Certifications
add_heading("Certifications", level=1)
add_paragraph(
    "- SQA Testing • Selenium with Java • API Testing with Postman • Manual Testing\n"
    "- Data Science (NLP, Cleaning, Visualization) • Apigee API Design"
)

# Save your document
doc.save("Samra_Darakshan_QA_Resume.docx")
