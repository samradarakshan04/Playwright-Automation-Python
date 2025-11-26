from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_paragraph(doc, text, bold=False, italic=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    font = run.font
    font.size = Pt(font_size)
    p.space_after = Pt(space_after)
    return p

def add_bullet_point(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

def create_resume():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header
    name = doc.add_paragraph()
    run = name.add_run('Samra Darakshan')
    run.bold = True
    run.font.size = Pt(16)
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    contacts = doc.add_paragraph()
    contacts.add_run('Email: samra.darakshan@gmail.com | Phone: +923352291683\n').bold = False
    contacts.add_run('LinkedIn: ').bold = True
    contacts.add_run('https://linkedin.com/in/samra-darakshan\n')
    contacts.add_run('GitHub: ').bold = True
    contacts.add_run('qa-testing-portfolio-1 | qa-testing-portfolio-2\n')
    contacts.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contacts.space_after = Pt(12)

    # Professional Summary
    add_heading(doc, 'Professional Summary', level=2)
    add_paragraph(doc,
    "Experienced Software Quality Assurance Engineer and Automation Engineer with over 7 years of expertise in manual and automated testing across web, mobile, and chatbot platforms. Skilled in building and maintaining robust test automation frameworks using Selenium (Java/Python), Playwright, TestNG, PyTest, Cypress, and Appium. Strong hands-on experience with API testing using Postman and Python, and integrating test automation into CI/CD pipelines with Jenkins and Git-based version control.\n"
    "Passionate about improving software quality through automation and manual testing, optimizing workflows, and enhancing user experience. Lead UX auditing initiatives for post-release reviews of customer and driver apps, and web portals to ensure consistency, usability, and polished user interfaces. Experienced in Agile/Scrum environments, delivering high-quality products that meet user expectations and business goals."
    )

    # Skills
    add_heading(doc, 'Skills', level=2)
    skills_text = (
        "Automation Testing: Selenium (Java/Python), Appium, PyTest, TestNG, JUnit, Cucumber, Cypress, Playwright\n"
        "API Testing: Postman (JavaScript & Python), API Automation, API Chaining, Chai Assertions, Newman\n"
        "Programming Languages: Python, Java, JavaScript, SQL\n"
        "Testing Methodologies: Agile, BDD, Manual Testing, Test Planning, Test Case Creation, Test Execution, Defect Reporting, Test Reporting, STLC\n"
        "Version Control: Git, GitHub, Bitbucket, TFS\n"
        "CI/CD Tools: Jenkins, Docker\n"
        "Soft Skills: Collaboration, Communication, Problem-solving, Time Management"
    )
    add_paragraph(doc, skills_text)

    # Professional Experience
    add_heading(doc, 'Professional Experience', level=2)

    # eZhire Role
    add_paragraph(doc, 'Software Quality Assurance Engineer | Automation Engineer', bold=True)
    add_paragraph(doc, 'eZhire, Karachi, Pakistan | Nov 2024 – Present', italic=True)
    bullets = [
        "Automated mobile app APIs for booking creation (daily, weekly, monthly) with various payment methods using Postman (JavaScript).",
        "Developed a scalable automation framework using Selenium with Python (PyTest) following Page Object Model (POM) for vendor assignments, car, and driver rental workflows.",
        "Automated key test scenarios for the eZhire customer app on Android and iOS using Appium and Python, increasing test coverage and reducing manual efforts.",
        "Led comprehensive manual testing of customer app, driver app, and admin dashboard to ensure smooth user experience and critical bug identification.",
        "Spearheaded post-release UX audits for eZhire’s customer app, driver app, and web admin portal by reviewing live UI screens, notifications, chatbot flows, emails, and error messages to improve consistency, usability, and branding alignment.",
        "Collaborated with Agile teams and used Git, GitHub, and Bitbucket for version control and Jenkins for CI/CD pipeline integration to maintain quality and delivery efficiency."
    ]
    for bullet in bullets:
        add_bullet_point(doc, bullet)

    # Royal Cyber Inc.
    add_paragraph(doc, 'SQA Engineer | Automation Engineer', bold=True)
    add_paragraph(doc, 'Royal Cyber Inc., Karachi, Pakistan | Mar 2022 – Oct 2024', italic=True)
    bullets = [
        "Built automation frameworks using Selenium (Java), Postman (API automation), integrated with Jenkins for CI/CD pipelines.",
        "Automated front-end web testing with Cypress and load testing with JMeter.",
        "Developed API automation scripts using Postman with API chaining, Chai assertions, and environment variables to validate complex workflows.",
        "Automated Admin APIs for IBM Integration Bus (IIB) and MQ, integrated with IBM Watson Assistant and Salesforce Einstein Copilot using Generative AI technologies.",
        "Performed black-box testing, regression, functional, and database testing using Jira for defect management and test tracking."
    ]
    for bullet in bullets:
        add_bullet_point(doc, bullet)

    # Systems Limited
    add_paragraph(doc, 'Software Engineer', bold=True)
    add_paragraph(doc, 'Systems Limited, Karachi, Pakistan | May 2020 – Mar 2022', italic=True)
    bullets = [
        "Developed management portals using Angular and Strapi CMS.",
        "Conducted manual testing to validate usability and functionality across platforms.",
        "Worked on REST API development and testing integrated with IIB, ACE, and MQ.",
        "Developed backend services using Java, Angular, and SQL."
    ]
    for bullet in bullets:
        add_bullet_point(doc, bullet)

    # Royal Cyber Inc. earlier role
    add_paragraph(doc, 'Software Engineer', bold=True)
    add_paragraph(doc, 'Royal Cyber Inc., Karachi, Pakistan | Apr 2018 – May 2020', italic=True)
    bullets = [
        "Developed React forms for a migration project with Kafka messaging integration.",
        "Specialized in MuleSoft API development and OpenAPI (Swagger) specification.",
        "Collaborated with QA teams on API testing best practices.",
        "Implemented DevOps solutions including Docker containers and Jenkins pipelines for cloud migration projects."
    ]
    for bullet in bullets:
        add_bullet_point(doc, bullet)

    # QA Intern
    add_paragraph(doc, 'QA Intern', bold=True)
    add_paragraph(doc, 'Sui Southern Gas Company (SSGC), Karachi, Pakistan | May 2017 – Aug 2017', italic=True)
    bullets = [
        "Performed manual testing on API, database, and web applications.",
        "Created test cases, reported defects, and participated in issue resolution meetings."
    ]
    for bullet in bullets:
        add_bullet_point(doc, bullet)

    # Education
    add_heading(doc, 'Education', level=2)
    add_paragraph(doc, 'MS in Data Science', bold=True)
    add_paragraph(doc, 'National University of Computing and Emerging Sciences (FAST), Karachi, Pakistan | Aug 2019 – Aug 2021', italic=True)
    add_paragraph(doc, 'CGPA: 3.53')

    add_paragraph(doc, 'BS in Computer Science', bold=True)
    add_paragraph(doc, 'Bahria University, Karachi, Pakistan | Feb 2014 – Mar 2018', italic=True)
    add_paragraph(doc, 'CGPA: 3.36')

    # Certifications
    add_heading(doc, 'Certifications', level=2)
    certs = [
        "Software Quality Assurance Certification",
        "Web Automation using Selenium with Java",
        "API Automation using Postman",
        "Manual Testing",
        "Introduction to Data Science | Cleaning Data in Python",
        "Supervised Learning | Data Visualization | Introduction to NLP",
        "API Design and Fundamentals of Google Cloud's Apigee API Platform"
    ]
    for cert in certs:
        add_bullet_point(doc, cert)

    # Vlogs Section Added Here
    add_heading(doc, 'Vlogs', level=2)
    vlogs = [
        "Exploring the Impact of Generative AI on APIs",
        "How API Connect and AI Transform Business",
        "The Migration Journey from IIB to ACE Simplified"
    ]
    for vlog in vlogs:
        add_bullet_point(doc, vlog)

    # Professional Projects
    add_heading(doc, 'Professional Projects', level=2)

    # Automation & Testing Projects
    add_paragraph(doc, 'Automation & Testing Projects', bold=True)
    proj_bullets = [
        "eZhire Mobile App API Automation: Automated booking APIs with multiple payment options using Postman and JavaScript.",
        "Vendor Assignments & Rentals Automation: Created automation framework using Selenium and Python (PyTest) to automate vendor and rental workflows.",
        "Driver App API Automation: Automated driver app workflows using Postman.",
        "Mobile Testing (eZhire Customer App): Automated Android/iOS mobile testing with Appium to increase coverage and efficiency.",
        "CI/CD & Version Control: Integrated Git/GitHub repositories and Jenkins pipelines to automate testing and deployment.",
        "Team Management Portal: Developed automation framework with Selenium, Java, TestNG, and Cucumber.",
        "Trello API Automation: Automated Trello API testing using Postman.",
        "OrangeHRM Automation: Automated tests using Selenium.",
        "Cypress POC: Implemented Cypress for cross-browser web app testing.",
        "Generative AI Integrations: Enhanced IBM Watson Assistant and Salesforce Einstein with generative AI capabilities for better middleware communication."
    ]
    for bullet in proj_bullets:
        add_bullet_point(doc, bullet)

    # API Integration & UI Development
    add_paragraph(doc, 'API Integration & UI Development', bold=True)
    api_ui_bullets = [
        "Integrated MuleSoft APIs with ELK, Salesforce, and BigQuery for Mattress Firm project.",
        "Developed React web apps for utility migration projects."
    ]
    for bullet in api_ui_bullets:
        add_bullet_point(doc, bullet)

    # Banking & API Management
    add_paragraph(doc, 'Banking & API Management', bold=True)
    bank_bullets = [
        "Worked on Apigee, RapidAPI, and Swagger API management platforms at Paccar.",
        "Developed and tested REST APIs for Meezan Bank, Bank Islami, and Dubai Islami Bank.",
        "Implemented IIB, SQL, and MQ integrations for Etisalat supply chain management."
    ]
    for bullet in bank_bullets:
        add_bullet_point(doc, bullet)

    # Save document
    doc.save('Samra_Darakshan_Resume.docx')
    print("Resume saved as 'Samra_Darakshan_Resume.docx'")

if __name__ == '__main__':
    create_resume()
